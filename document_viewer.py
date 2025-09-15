import streamlit as st
import PyPDF2
import docx
import base64
import io
from io import BytesIO

def display_pdf(file_data, display_name):
    """Display a PDF file using Streamlit."""
    # Create a temporary in-memory file
    pdf_bytes = BytesIO(file_data)
    
    # Container for the PDF viewer with styling
    with st.container(border=True):
        # PDF header with title and download button
        col1, col2 = st.columns([3, 1])
        with col1:
            st.subheader(f"📕 {display_name}")
        with col2:
            # Create a download button for the PDF
            st.download_button(
                label="⬇️ Download",
                data=pdf_bytes,
                file_name=display_name,
                mime="application/pdf",
                use_container_width=True
            )
        
        # Document info
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_bytes)
            num_pages = len(pdf_reader.pages)
            st.caption(f"PDF Document • {num_pages} pages")
        except:
            st.caption("PDF Document")
        
        # Convert PDF bytes to base64 for embedding
        pdf_bytes.seek(0)  # Reset the BytesIO pointer
        base64_pdf = base64.b64encode(pdf_bytes.read()).decode('utf-8')
        
        # Create PDF viewer with custom styling
        pdf_display = f"""
        <style>
        .pdf-viewer {{
            width: 100%;
            height: 800px;
            border: 1px solid #e0e0e0;
            border-radius: 5px;
        }}
        </style>
        <iframe class="pdf-viewer" src="data:application/pdf;base64,{base64_pdf}" type="application/pdf"></iframe>
        """
        st.markdown(pdf_display, unsafe_allow_html=True)

def display_docx(file_data, display_name):
    """Display a DOCX file in Streamlit with enhanced UI."""
    # Create a temporary in-memory file
    docx_bytes = BytesIO(file_data)
    
    with st.container(border=True):
        # Document header with title and download button
        col1, col2 = st.columns([3, 1])
        with col1:
            st.subheader(f"📘 {display_name}")
        with col2:
            # Create a download button for the DOCX
            st.download_button(
                label="⬇️ Download",
                data=docx_bytes,
                file_name=display_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        try:
            # Open the document
            doc = docx.Document(docx_bytes)
            
            # Document info
            st.caption(f"Word Document • {len(doc.paragraphs)} paragraphs • {len(doc.tables)} tables")
            
            # Display content in a scrollable container
            with st.container(height=700, border=False):
                # Display document with styling
                st.markdown("""
                <style>
                .docx-content {
                    background-color: grey;
                    padding: 20px;
                    border-radius: 5px;
                    border: 1px solid #e0e0e0;
                }
                .docx-heading {
                    margin-top: 20px;
                    margin-bottom: 10px;
                }
                .docx-paragraph {
                    margin-bottom: 10px;
                    line-height: 1.5;
                }
                </style>
                """, unsafe_allow_html=True)
                
                st.markdown('<div class="docx-content">', unsafe_allow_html=True)
                
                # Display paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():  # Skip empty paragraphs
                        # Check if this is a heading
                        if para.style.name.startswith('Heading'):
                            # Use appropriate Markdown for headings
                            heading_level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                            st.markdown(f'<h{heading_level} class="docx-heading">{para.text}</h{heading_level}>', unsafe_allow_html=True)
                        else:
                            st.markdown(f'<p class="docx-paragraph">{para.text}</p>', unsafe_allow_html=True)
                
                # Display tables if any
                for i, table in enumerate(doc.tables):
                    st.markdown(f'<h3 class="docx-heading">Table {i+1}</h3>', unsafe_allow_html=True)
                    # Convert table to HTML
                    html_table = "<table style='width:100%; border-collapse: collapse; margin-bottom: 20px;'>"
                    for j, row in enumerate(table.rows):
                        html_table += "<tr>"
                        for cell in row.cells:
                            # Add header styling if this is the first row
                            if j == 0:
                                html_table += f"<th style='border: 1px solid #ddd; padding: 8px; background-color: #f2f2f2;'>{cell.text}</th>"
                            else:
                                html_table += f"<td style='border: 1px solid #ddd; padding: 8px;'>{cell.text}</td>"
                        html_table += "</tr>"
                    html_table += "</table>"
                    st.markdown(html_table, unsafe_allow_html=True)
                
                st.markdown('</div>', unsafe_allow_html=True)
                
        except Exception as e:
            st.error(f"Error displaying DOCX file: {str(e)}")
            st.info("You can download the file to view it in Microsoft Word or another compatible application.")

def display_text(file_data, display_name):
    """Display a text file in Streamlit with enhanced UI."""
    with st.container(border=True):
        # Document header with title and download button
        col1, col2 = st.columns([3, 1])
        with col1:
            st.subheader(f"📄 {display_name}")
        with col2:
            # Create a download button for the text file
            st.download_button(
                label="⬇️ Download",
                data=file_data,
                file_name=display_name,
                mime="text/plain",
                use_container_width=True
            )
        
        # Display the text content
        try:
            text_content = file_data.decode('utf-8')
            # Get line count for info
            line_count = text_content.count('\n') + 1
            char_count = len(text_content)
            st.caption(f"Text Document • {line_count} lines • {char_count} characters")
            
            # Display in a styled text area
            st.markdown("""
            <style>
            .text-content {
                font-family: monospace;
                white-space: pre-wrap;
                background-color: #f8f9fa;
                padding: 15px;
                border-radius: 5px;
                border: 1px solid #e0e0e0;
                height: 600px;
                overflow-y: auto;
            }
            </style>
            """, unsafe_allow_html=True)
            
            st.markdown(f'<div class="text-content">{text_content}</div>', unsafe_allow_html=True)
        except UnicodeDecodeError:
            # Try with a different encoding if UTF-8 fails
            try:
                text_content = file_data.decode('latin-1')
                st.text_area("File Content", text_content, height=600)
            except Exception as e:
                st.error(f"Error displaying text file: {str(e)}")

def display_document(file_data, file_info):
    """Display a document based on its file type with enhanced UI."""
    file_type = file_info.get("file_type", "unknown")
    display_name = file_info.get("display_name", file_info.get("filename", "Unnamed Document"))
    
    st.title("Document Viewer")
    
    # Document metadata
    col1, col2, col3 = st.columns([2, 1, 1])
    with col1:
        st.write(f"**Name:** {display_name}")
    with col2:
        st.write(f"**Type:** {file_type.upper()}")
    with col3:
        if "upload_date" in file_info:
            st.write(f"**Uploaded:** {file_info['upload_date'].strftime('%Y-%m-%d')}")
    
    # Divider
    st.divider()
    
    # Display document based on file type
    if file_type == "pdf":
        display_pdf(file_data, display_name)
    elif file_type in ["docx", "doc"]:
        display_docx(file_data, display_name)
    elif file_type == "txt":
        display_text(file_data, display_name)
    else:
        st.error(f"Unsupported file type: {file_type}")
        # Still provide a download option
        st.download_button(
            label=f"Download {file_type.upper()} File",
            data=io.BytesIO(file_data),
            file_name=display_name,
        )