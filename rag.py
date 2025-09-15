import os
import tempfile
import PyPDF2
import streamlit as st
import torch
import requests
from bs4 import BeautifulSoup
import time
import pandas as pd
import warnings
import threading
import psutil
import io
import docx
from utils import remove_directory_recursively
import datetime
import bson
import hashlib
import json

# LangChain imports
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import OllamaLLM
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS, Chroma
from langchain.chains import RetrievalQA, LLMChain
from langchain.docstore.document import Document
from langchain.prompts import PromptTemplate
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler

# Suppress warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)
os.environ["PYTHONWARNINGS"] = "ignore::DeprecationWarning"

# Fix for PyTorch/Streamlit compatibility issue
if "STREAMLIT_WATCH_MODULES" in os.environ:
    modules_to_skip = ["torch", "tensorflow"]
    current_modules = os.environ["STREAMLIT_WATCH_MODULES"].split(",")
    filtered_modules = [m for m in current_modules if all(skip not in m for skip in modules_to_skip)]
    os.environ["STREAMLIT_WATCH_MODULES"] = ",".join(filtered_modules)

class EnhancedRAG:
    def __init__(self, 
                 llm_model_name="llama3.2:latest",
                 embedding_model_name="sentence-transformers/all-MiniLM-L6-v2",
                 chunk_size=1000,
                 chunk_overlap=200,
                 use_gpu=True):
        """
        Initialize the Enhanced RAG system with multiple modes.
        
        Args:
            llm_model_name: The Ollama model for text generation
            embedding_model_name: The HuggingFace model for embeddings
            chunk_size: Size of document chunks
            chunk_overlap: Overlap between chunks
            use_gpu: Whether to use GPU acceleration
        """
        self.llm_model_name = llm_model_name
        self.embedding_model_name = embedding_model_name
        self.use_gpu = use_gpu and torch.cuda.is_available()
        self.temp_dirs = []  # Keep track of temporary directories
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Device selection for embeddings
        self.device = "cuda" if self.use_gpu else "cpu"
        st.sidebar.info(f"Using device: {self.device}")
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
        
        # Initialize embeddings model
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name=embedding_model_name,
                model_kwargs={"device": self.device}
            )
            # Store embedding function for easier access
            self.embedding_function = self.embeddings
            st.sidebar.success(f"Embeddings model loaded: {embedding_model_name}")
        except Exception as e:
            st.sidebar.error(f"Failed to load embeddings model: {str(e)}")
            self.embeddings = None
        
        # Initialize LLM
        try:
            callbacks = [StreamingStdOutCallbackHandler()]
            self.llm = OllamaLLM(model=llm_model_name, callbacks=callbacks)
            st.sidebar.success(f"LLM loaded: {llm_model_name}")
        except Exception as e:
            st.sidebar.error(f"Failed to load LLM: {str(e)}")
            self.llm = None
        
        # Initialize vector stores for different sources
        self.doc_vector_store = None
        self.web_vector_store = None
        self.documents_processed = 0
        self.documents = []
        
        # Monitoring stats
        self.processing_times = {}
        
        # Keep track of sources and errors
        self.sources = []
        self.errors = []
        
        # Blockchain support
        self.use_blockchain = False
        self.blockchain = None
    
    
    def __del__(self):
        """Cleanup temporary directories when object is garbage collected."""
        for temp_dir in self.temp_dirs:
            try:
                if os.path.exists(temp_dir):
                    remove_directory_recursively(temp_dir)
            except:
                pass

    def initialize_blockchain(self, blockchain_url, contract_address, private_key, chain_id=1337):
        """Initialize blockchain integration for document verification and query logging.
        
        Args:
            blockchain_url: URL of the blockchain node
            contract_address: Address of the deployed RAG Document Verifier contract
            private_key: Private key for signing transactions
            chain_id: Chain ID of the blockchain network
            
        Returns:
            bool: Success status
        """
        try:
            # First, print detailed debug information
            print(f"Initializing blockchain with:")
            print(f"  - URL: {blockchain_url}")
            print(f"  - Contract: {contract_address}")
            print(f"  - Chain ID: {chain_id}")
            
            # First, check if all required parameters are provided
            if not blockchain_url:
                st.warning("⚠️ No blockchain URL provided. System will operate without blockchain verification.")
                self.use_blockchain = False
                return False
                
            if not contract_address:
                st.warning("⚠️ No contract address provided. System will operate without blockchain verification.")
                self.use_blockchain = False
                return False
                
            if not private_key:
                st.warning("⚠️ No private key provided. System will operate without blockchain verification.")
                self.use_blockchain = False
                return False
            
            # Import the BlockchainManager
            try:
                from blockchain_utils import BlockchainManager
                print(f"Successfully imported BlockchainManager")
            except ImportError as ie:
                st.error(f"❌ Blockchain module not available: {str(ie)}. Make sure blockchain_utils.py is in your project.")
                self.use_blockchain = False
                return False
            
            # Create the blockchain manager
            try:
                print(f"Creating BlockchainManager instance...")
                self.blockchain = BlockchainManager(
                    blockchain_url=blockchain_url,
                    contract_address=contract_address,
                    private_key=private_key,
                    chain_id=chain_id
                )
                
                # Test the connection and contract
                print(f"Testing blockchain connection...")
                connection_test = self.blockchain.test_connection()
                
                # Print detailed connection test results
                print(f"Connection test results: {connection_test}")
                
                # Check if we're in simulation mode
                if self.blockchain.simulation_mode:
                    st.warning(f"⚠️ Running in blockchain simulation mode. Your data will NOT be stored on a real blockchain.")
                    # Provide a more detailed reason if available from test results
                    reason = self.blockchain.connection_error
                    if not reason and connection_test.get("errors"):
                        reason = connection_test["errors"][-1] # Get the last error from the test connection
                    st.info(f"Reason: {reason}")
                    st.info("All blockchain operations will be simulated.")
                    self.use_blockchain = True
                    return True
                
                if connection_test.get("connection", False):
                    if not connection_test.get("account", False):
                        st.warning(f"⚠️ Connected to blockchain, but account could not be loaded. Check private key.")
                        self.use_blockchain = False
                        return False
                    
                    if not connection_test.get("contract", False):
                        # Try deploying the contract automatically
                        st.info("📄 Contract not found or not accessible. Attempting to deploy contract automatically...")
                        
                        # Deploy the contract via BlockchainManager
                        if hasattr(self.blockchain, 'deploy_contract') and callable(self.blockchain.deploy_contract):
                            try:
                                print(f"Attempting to deploy contract...")
                                deploy_result = self.blockchain.deploy_contract(self.blockchain.contract.abi)
                                if deploy_result:
                                    st.success("✅ Contract deployed successfully!")
                                    # Update test results
                                    connection_test = self.blockchain.test_connection()
                                else:
                                    st.error("❌ Failed to deploy contract. Please check blockchain connection and account.")
                                    self.use_blockchain = False
                                    return False
                            except Exception as deploy_error:
                                st.error(f"❌ Error deploying contract: {str(deploy_error)}")
                                self.use_blockchain = False
                                return False
                        else:
                            st.error("❌ Contract deployment function not available.")
                            self.use_blockchain = False
                            return False
                    
                    if not connection_test.get("read_test", False):
                        errors = connection_test.get("errors", ["Unknown error"])
                        st.warning(f"⚠️ Contract methods not working: {', '.join(errors)}")
                        self.use_blockchain = False
                        return False
                    
                    # Everything looks good!
                    self.use_blockchain = True
                    print(f"Blockchain verification activated successfully!")
                    st.success("✅ Blockchain verification activated successfully!")
                    return True
                else:
                    errors = connection_test.get("errors", ["Unknown error"])
                    error_msg = f"⚠️ Could not connect to blockchain: {', '.join(errors)}"
                    print(error_msg)
                    st.warning(error_msg)
                    self.use_blockchain = False
                    return False
                    
            except Exception as e:
                import traceback
                traceback.print_exc()
                st.error(f"❌ Error initializing blockchain: {str(e)}")
                self.use_blockchain = False
                return False
                
        except Exception as e:
            import traceback
            traceback.print_exc()
            st.error(f"❌ Unexpected error with blockchain: {str(e)}")
            self.use_blockchain = False
            return False

    def verify_document_blockchain(self, file_data, document_id):
        """Verify a document on the blockchain.
        
        Args:
            file_data: Binary content of the document
            document_id: Unique identifier for the document
            
        Returns:
            dict: Verification data or None if verification failed
        """
        if not hasattr(self, 'use_blockchain') or not self.use_blockchain or not self.blockchain:
            print("Blockchain verification skipped - blockchain not enabled or initialized")
            return None
            
        try:
            import tempfile
            import os
            
            # Create a temporary file for verification
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                temp_file.write(file_data)
                temp_path = temp_file.name
                
            # Verify the document
            verification = self.blockchain.verify_document(document_id, temp_path)
            
            # Clean up temporary file
            try:
                os.unlink(temp_path)
            except:
                pass
                
            if verification and verification.get('status') == 1:  # Success
                result = {
                    "verified": True,
                    "document_id": document_id,
                    "document_hash": verification.get("document_hash", ""),
                    "tx_hash": verification.get("tx_hash", ""),
                    "block_number": verification.get("block_number", 0),
                    "timestamp": datetime.datetime.now().isoformat()
                }
                
                # Add simulation flag if in simulation mode
                if hasattr(self.blockchain, 'simulation_mode') and self.blockchain.simulation_mode:
                    result["simulation"] = True
                
                return result
            return None
        except Exception as e:
            print(f"Blockchain verification error: {str(e)}")
            return None

    def log_query_blockchain(self, query, answer):
        """Log a query and its answer on the blockchain.
        
        Args:
            query: The user's query text
            answer: The generated answer text
            
        Returns:
            dict: Logging result or None if logging failed
        """
        if not hasattr(self, 'use_blockchain') or not self.use_blockchain or not self.blockchain:
            print("Query logging to blockchain skipped - blockchain not enabled or initialized")
            return None
            
        try:
            # Log query to blockchain
            log_result = self.blockchain.log_query(query, answer)
            
            # Check if we're in simulation mode
            simulation_mode = hasattr(self.blockchain, 'simulation_mode') and self.blockchain.simulation_mode
            
            if log_result and log_result.get("status") == 1:  # Success
                result = {
                    "logged": True,
                    "query_id": log_result.get("query_id", ""),
                    "tx_hash": log_result.get("tx_hash", ""),
                    "block_number": log_result.get("block_number", 0),
                    "timestamp": datetime.datetime.now().isoformat()
                }
                
                # Add simulation flag if in simulation mode
                if simulation_mode:
                    result["simulation"] = True
                    
                return result
            return None
        except Exception as e:
            print(f"Error logging query to blockchain: {str(e)}")
            return None

    def process_files(self, files, user_id=None, mongodb=None, notebook_id=None, is_nested=False, domains=None):
        """Process files and build vector store.
        
        Args:
            files: List of file objects
            user_id: Optional user ID for logging
            mongodb: Optional MongoDB connection
            notebook_id: Optional notebook ID to associate with documents
            is_nested: Whether this is being called from within another streamlit component
            domains: Optional list of domains/topics this data belongs to
            
        Returns:
            Boolean indicating success
        """
        # Check if embeddings are available
        if self.embeddings is None:
            st.error("Embeddings model not initialized. Unable to process files.")
            return False
            
        all_docs = []
        document_metadata = []
        
        # Use status or simple spinner based on nesting
        if is_nested:
            # We're inside another expander/status, so just use progress indicators
            status_msg = st.empty()
            status_msg.info("Processing files...")
            progress_bar = st.progress(0)
        else:
            # We can use the full status widget
            status = st.status("Processing files...")
        
        # Create temporary directory for file storage
        temp_dir = tempfile.mkdtemp()
        self.temp_dirs.append(temp_dir)  # Track for cleanup
        st.session_state['temp_dir'] = temp_dir
        
        # Monitor processing time and memory usage
        start_time = time.time()
        
        # Track memory before processing
        mem_before = psutil.virtual_memory().used / (1024 * 1024 * 1024)  # GB
        
        # Process each file
        total_files = len(files)
        for i, file in enumerate(files):
            try:
                # Update progress
                if is_nested:
                    progress_bar.progress((i + 1) / total_files)
                    status_msg.info(f"Processing {file.name} ({i+1}/{total_files})...")
                else:
                    status.update(label=f"Processing {file.name} ({i+1}/{total_files})...")
                
                file_start_time = time.time()
                file_type = "unknown"
                
                # Determine file type
                if file.name.lower().endswith('.pdf'):
                    file_type = "pdf"
                elif file.name.lower().endswith(('.docx', '.doc')):
                    file_type = "docx"
                elif file.name.lower().endswith(('.txt')):
                    file_type = "txt"
                
                # Save uploaded file to temp directory
                file_path = os.path.join(temp_dir, file.name)
                file.seek(0)  # Reset file pointer
                file_content = file.read()  # Read the file content
                
                with open(file_path, "wb") as f:
                    f.write(file_content)
                
                # Extract text based on file type
                text = ""
                page_count = 0
                
                if file_type == "pdf":
                    try:
                        with open(file_path, "rb") as f:
                            pdf = PyPDF2.PdfReader(f)
                            page_count = len(pdf.pages)
                            for page_num in range(page_count):
                                page = pdf.pages[page_num]
                                page_text = page.extract_text()
                                if page_text:
                                    text += page_text + "\n\n"
                    except Exception as e:
                        st.error(f"Error extracting text from PDF {file.name}: {str(e)}")
                        continue
                
                elif file_type == "docx":
                    try:
                        doc = docx.Document(file_path)
                        page_count = len(doc.paragraphs)
                        for para in doc.paragraphs:
                            text += para.text + "\n\n"
                    except Exception as e:
                        st.error(f"Error extracting text from DOCX {file.name}: {str(e)}")
                        continue
                
                elif file_type == "txt":
                    try:
                        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                            text = f.read()
                        page_count = text.count("\n") + 1
                    except Exception as e:
                        st.error(f"Error extracting text from TXT {file.name}: {str(e)}")
                        continue
                
                # Skip file if no text was extracted
                if not text.strip():
                    st.warning(f"No text content found in {file.name}. Skipping...")
                    continue
                
                # Create documents
                docs = [Document(page_content=text, metadata={
                    "source": file.name, 
                    "notebook_id": notebook_id,
                    "file_type": file_type
                })]
                
                # Split documents into chunks
                split_docs = self.text_splitter.split_documents(docs)
                
                all_docs.extend(split_docs)
                
                file_end_time = time.time()
                processing_time = file_end_time - file_start_time
                
                # Save document metadata
                doc_meta = {
                    "filename": file.name,
                    "file_type": file_type,
                    "page_count": page_count,
                    "chunk_count": len(split_docs),
                    "processing_time": processing_time,
                    "notebook_id": notebook_id
                }
                document_metadata.append(doc_meta)
                
                # Blockchain verification (if enabled)
                if self.use_blockchain and self.blockchain:
                    try:
                        # Create a unique document ID
                        document_id = f"{file.name}_{notebook_id}_{datetime.datetime.now().timestamp()}"
                        
                        # Verify document on blockchain
                        if is_nested:
                            status_msg.info(f"Verifying {file.name} on blockchain...")
                        else:
                            status.update(label=f"Verifying {file.name} on blockchain...")
                            
                        verification = self.verify_document_blockchain(file_content, document_id)
                        
                        if verification:
                            # Add blockchain verification metadata to document
                            doc_meta["blockchain_verification"] = verification
                            
                            # Add blockchain verification to all chunks
                            for doc in split_docs:
                                doc.metadata["blockchain_verification"] = verification
                                
                            # Success message
                            if is_nested:
                                st.success(f"✅ {file.name} verified on blockchain")
                            else:
                                st.sidebar.success(f"✅ {file.name} verified on blockchain")
                                st.sidebar.info(f"Transaction: {verification['tx_hash'][:10]}...")
                    except Exception as e:
                        error_msg = f"Blockchain verification error: {str(e)}"
                        if is_nested:
                            st.error(error_msg)
                        else:
                            st.sidebar.error(error_msg)
                
                # Display success message
                success_msg = f"Processed {file.name}: {len(split_docs)} chunks in {processing_time:.2f}s"
                if is_nested:
                    st.success(success_msg)
                else:
                    st.sidebar.success(success_msg)
                    
                self.processing_times[file.name] = {
                    "chunks": len(split_docs),
                    "time": processing_time
                }
                
            except Exception as e:
                error_msg = f"Error processing {file.name}: {str(e)}"
                self.errors.append(error_msg)
                if is_nested:
                    st.error(error_msg)
                else:
                    st.sidebar.error(error_msg)
        
        # Create vector store if we have documents
        if all_docs:
            if is_nested:
                status_msg.info("Building vector index...")
            else:
                status.update(label="Building vector index...")
                
            try:
                # Record the time taken to build the index
                index_start_time = time.time()
                
                # Create the vector store using FAISS
                self.doc_vector_store = FAISS.from_documents(all_docs, self.embeddings)
                
                # Store all documents for later use
                self.documents = all_docs
                
                index_end_time = time.time()
                index_time = index_end_time - index_start_time
                
                # Track memory after processing
                mem_after = psutil.virtual_memory().used / (1024 * 1024 * 1024)  # GB
                mem_used = mem_after - mem_before
                
                total_time = time.time() - start_time
                
                complete_msg = f"Completed processing {len(all_docs)} chunks in {total_time:.2f}s"
                if is_nested:
                    status_msg.success(complete_msg)
                    progress_bar.progress(1.0)
                else:
                    status.update(label=complete_msg, state="complete")
                
                # Save performance metrics
                self.processing_times["index_building"] = index_time
                self.processing_times["total_time"] = total_time
                self.processing_times["memory_used_gb"] = mem_used
                self.documents_processed = len(all_docs)
                
                # Save document metadata to MongoDB if user is logged in
                if user_id and mongodb:
                    overall_meta = {
                        "documents": document_metadata,
                        "total_chunks": len(all_docs),
                        "index_building_time": index_time,
                        "total_processing_time": total_time,
                        "memory_used_gb": mem_used,
                        "notebook_id": notebook_id
                    }
                    mongodb.save_document_metadata(user_id, overall_meta, notebook_id)
                
                # Clean up temporary UI elements if nested
                if is_nested:
                    time.sleep(1)  # Let user see completion message
                    status_msg.empty()
                    # Keep progress bar at 100%
                
                # Auto-detect domains if not provided
                if not domains and len(self.documents) > 0:
                    domains = self.detect_domains(self.documents)
                
                # If MongoDB and notebook_id are provided, save the vector index
                if mongodb and notebook_id:
                    try:
                        # Serialize the FAISS index
                        import faiss
                        import pickle
                        import base64
                        from io import BytesIO
                        
                        # Get the FAISS index from the vector store
                        faiss_index = self.doc_vector_store.index
                        
                        # Serialize the index
                        buffer = BytesIO()
                        faiss.write_index(faiss_index, buffer)
                        serialized_index = base64.b64encode(buffer.getvalue()).decode()
                        
                        # Prepare metadata
                        metadata = {
                            "embedding_model": self.embedding_model_name,
                            "chunk_size": self.chunk_size,
                            "chunk_overlap": self.chunk_overlap,
                            "document_count": len(self.documents),
                            "index_size_bytes": len(serialized_index),
                            "domains": domains or []  # Add detected domains
                        }
                        
                        # Also serialize document info
                        doc_buffer = BytesIO()
                        pickle.dump(self.documents, doc_buffer)
                        serialized_docs = base64.b64encode(doc_buffer.getvalue()).decode()
                        
                        # Combine both serialized objects
                        index_data = {
                            "faiss_index": serialized_index,
                            "documents": serialized_docs
                        }
                        
                        # Save to MongoDB
                        mongodb.save_faiss_index(notebook_id, user_id, index_data, metadata)
                        
                        if not is_nested:
                            st.success(f"Vector index saved for {', '.join(domains) if domains else 'your documents'}")
                            
                    except Exception as e:
                        if not is_nested:
                            st.error(f"Failed to save vector index: {str(e)}")
                
                # Save the vector store if mongodb and notebook_id are provided
                if mongodb and notebook_id and hasattr(self, 'doc_vector_store') and self.doc_vector_store:
                    self.save_vector_store(mongodb, notebook_id, user_id, is_nested)
                
                # Add debugging
                try:
                    print("\n----- DEBUG INFO -----")
                    self.debug_vector_store()
                    print("----- END DEBUG INFO -----\n")
                except Exception as e:
                    print(f"Debug error: {str(e)}")
                
                return True
            except Exception as e:
                error_msg = f"Error creating vector store: {str(e)}"
                self.errors.append(error_msg)
                st.error(error_msg)
                if is_nested:
                    status_msg.error(error_msg)
                    time.sleep(2)  # Let user see error message
                    status_msg.empty()
                else:
                    status.update(label=error_msg, state="error")
                return False
        else:
            empty_msg = "No content extracted from files"
            if is_nested:
                status_msg.error(empty_msg)
                time.sleep(2)  # Let user see error message
                status_msg.empty()
            else:
                status.update(label=empty_msg, state="error")
            return False

    def detect_domains(self, documents, max_domains=3):
        """Auto-detect domains/topics in the documents.
        
        Args:
            documents: List of document objects
            max_domains: Maximum number of domains to detect
            
        Returns:
            List of domain strings
        """
        try:
            # Simple approach: extract key topics based on frequency
            from collections import Counter
            import re
            
            # Combine all text
            all_text = " ".join([doc.page_content for doc in documents])
            
            # Get potential domain keywords (simplified approach)
            common_domains = {
                "machine learning": ["algorithm", "model", "training", "neural", "dataset", "features"],
                "data science": ["analysis", "visualization", "statistics", "correlation", "hypothesis"],
                "programming": ["code", "function", "class", "variable", "algorithm", "programming"],
                "finance": ["market", "investment", "stock", "financial", "trading", "economy"],
                "healthcare": ["patient", "treatment", "medical", "clinical", "diagnosis", "health"]
            }
            
            # Count occurrences of domain keywords
            domain_scores = {}
            for domain, keywords in common_domains.items():
                score = sum(len(re.findall(r'\b' + re.escape(keyword) + r'\b', all_text.lower())) 
                           for keyword in keywords)
                domain_scores[domain] = score
            
            # Return top domains
            return [domain for domain, score in 
                   sorted(domain_scores.items(), key=lambda x: x[1], reverse=True)[:max_domains] 
                   if score > 0]
        except:
            # If anything fails, return empty list
            return []

    def enhance_answer(self, initial_answer, query, source_content):
        """
        Enhance the initial answer with additional context and improved quality.
        
        Args:
            initial_answer: The initial answer generated by the RAG system
            query: The original user query
            source_content: The source content chunks used to generate the answer
            
        Returns:
            An enhanced answer with improved quality and formatting
        """
        # Create an enhancement prompt template
        enhance_template = """
        You are an expert content enhancer. Your task is to improve the quality of an AI-generated answer
        while maintaining factual accuracy.
        
        Below is a query, an initial answer, and the source content used to generate that answer.
        
        QUERY:
        {query}
        
        INITIAL ANSWER:
        {initial_answer}
        
        SOURCE CONTENT (EXTRACT):
        {source_content}
        
        Please enhance the initial answer by:
        1. Improving clarity and readability
        2. Adding relevant details from the source if they were missed
        3. Ensuring all claims are factually supported by the source content
        4. Adding appropriate structure (headings, bullet points) if helpful
        5. Making sure the tone is professional and helpful
        6. if need exapmles for the topic try give a example for better understanding. 
        7. if you are having multiple sub topics in it try to ehance the matter in each subtopic for better quality.
        
        ENHANCED ANSWER:
        """
        
        # Create enhancement prompt
        enhancement_prompt = PromptTemplate(
            template=enhance_template,
            input_variables=["query", "initial_answer", "source_content"]
        )
        
        # Create enhancement chain
        enhancement_chain = LLMChain(
            llm=self.llm,
            prompt=enhancement_prompt
        )
        
        # Prepare source content for the enhancement (limited to avoid token limits)
        summarized_sources = "\n\n".join([
            f"SOURCE {i+1}:\n{source[:500]}..." if len(source) > 500 else f"SOURCE {i+1}:\n{source}"
            for i, source in enumerate(source_content[:3])  # Limit to first 3 sources
        ])
        
        # Invoke the enhancement chain
        try:
            enhanced_result = enhancement_chain.invoke({
                "query": query,
                "initial_answer": initial_answer,
                "source_content": summarized_sources
            })
            
            return enhanced_result["text"].strip()
        except Exception as e:
            st.warning(f"Enhancement step encountered an issue: {str(e)}. Using initial answer.")
            self.errors.append(f"Enhancement error: {str(e)}")
            return initial_answer

    def web_search(self, query, num_results=5):
        """
        Perform a web search using multiple fallback methods
        """
        try:
            # For this implementation, we'll use a simulated search
            # In a production environment, you would integrate with a real search API
            results = self.simulate_search(query, num_results)
            if results and len(results) > 0:
                self.errors.append("Search simulation succeeded")
                return results
            else:
                return self.get_mock_results(query)
        except Exception as e:
            self.errors.append(f"Search error: {str(e)}")
            return self.get_mock_results(query)

    def simulate_search(self, query, num_results=5):
        """Simulate web search results for a query.
        This provides plausible information even when no documents are available.
        
        Args:
            query: The user's query
            num_results: Number of search results to simulate
            
        Returns:
            List of search result dictionaries
        """
        # First see if we can get canned results for common topics
        canned_results = self.get_mock_results(query)
        if canned_results:
            return canned_results[:num_results]
        
        # Otherwise, generate simulated search results
        results = []
        
        # Generate plausible titles and snippets based on the query
        prompt = f"""
        Generate {num_results} plausible web search results for the query: "{query}"
        
        Each result should have:
        1. A realistic website name and URL
        2. A title that might appear in search results
        3. A brief snippet/content that might appear in search results (100-150 words)
        
        Format each result as:
        title: [TITLE]
        url: [URL]
        content: [CONTENT]
        
        Make the content informative and factually accurate.
        """
        
        try:
            response = self.llm(prompt)
            
            # Parse the response into separate results
            result_blocks = response.split("title:")[1:]  # Skip the text before the first result
            
            for block in result_blocks:
                if not block.strip():
                    continue
                
                # Extract components
                parts = block.split("url:", 1)
                if len(parts) < 2:
                    continue
                    
                title = parts[0].strip()
                
                remaining = parts[1].split("content:", 1)
                if len(remaining) < 2:
                    continue
                    
                url = remaining[0].strip()
                content = remaining[1].strip()
                
                results.append({
                    "title": title,
                    "url": url,
                    "content": content
                })
                
                if len(results) >= num_results:
                    break
                    
        except Exception as e:
            # Fallback if generation fails
            print(f"Error generating search results: {str(e)}")
            results = [
                {
                    "title": f"Informational resource about {query}",
                    "url": "https://example.com/info",
                    "content": f"This would contain information about {query}, but no external data is currently available. Please upload relevant documents for more specific information."
                }
            ]
        
        # Return at least one result
        return results if results else [
            {
                "title": f"Information about {query}",
                "url": "https://example.com/info",
                "content": f"Information about {query} would typically be found here. For more specific answers, consider uploading relevant documents."
            }
        ]

    def get_mock_results(self, query):
        """Get pre-defined search results for common topics.
        This provides more reliable information for common queries.
        
        Args:
            query: The user's query
            
        Returns:
            List of search result dictionaries or None if no matches
        """
        # Lowercase query for matching
        query_lower = query.lower()
        
        # Define common topic patterns and responses
        topics = {
            "machine learning": [
                {
                    "title": "Introduction to Machine Learning - MIT OpenCourseWare",
                    "url": "https://ocw.mit.edu/courses/electrical-engineering-and-computer-science/6-867-machine-learning-fall-2006/",
                    "content": "Machine learning is a field of computer science that gives computers the ability to learn without being explicitly programmed. Key components include: data preparation, algorithm selection, model training, evaluation, and deployment. Modern ML approaches include supervised learning, unsupervised learning, deep learning, and reinforcement learning."
                },
                {
                    "title": "Machine Learning Algorithms Explained - Towards Data Science",
                    "url": "https://towardsdatascience.com/machine-learning-algorithms-explained-3faf6cef544",
                    "content": "Machine learning algorithms include: Decision Trees, Random Forests, Support Vector Machines (SVMs), Neural Networks, k-Nearest Neighbors (k-NN), Linear/Logistic Regression, and clustering algorithms like K-means. Each algorithm has different strengths and is suitable for specific types of problems and data."
                },
                {
                    "title": "Getting Started with Machine Learning - Google Developers",
                    "url": "https://developers.google.com/machine-learning/guides/getting-started",
                    "content": "Building a machine learning system requires: collecting and preparing quality data, selecting appropriate features, choosing a suitable algorithm, training the model, evaluating performance using metrics like precision/recall or RMSE, and regular retraining to maintain accuracy over time."
                }
            ],
            "deep learning": [
                {
                    "title": "Deep Learning Explained - Stanford University",
                    "url": "https://cs.stanford.edu/people/karpathy/deeplearning/",
                    "content": "Deep learning is a subset of machine learning using neural networks with multiple layers. These deep neural networks can automatically learn hierarchical features from data. Key architectures include Convolutional Neural Networks (CNNs) for images, Recurrent Neural Networks (RNNs) for sequence data, and Transformers for NLP tasks."
                },
                {
                    "title": "Deep Learning vs. Machine Learning - IBM Research",
                    "url": "https://www.ibm.com/cloud/learn/deep-learning",
                    "content": "While traditional machine learning relies on structured data and feature engineering, deep learning can work with unstructured data and automatically extract features. Deep learning typically requires more data and computational resources but can achieve superior performance on complex tasks like image recognition and natural language processing."
                }
            ],
            "programming": [
                {
                    "title": "Learn Programming - freeCodeCamp",
                    "url": "https://www.freecodecamp.org/learn/",
                    "content": "Programming fundamentals include: variables, data types, control structures (if/else, loops), functions, and object-oriented concepts. Modern programming languages include Python, JavaScript, Java, C++, and Go. Good programming practices emphasize readability, modularity, testing, and documentation."
                },
                {
                    "title": "Programming Paradigms Explained - Medium",
                    "url": "https://medium.com/swlh/programming-paradigms-explained-simply-e8e8e7de078",
                    "content": "Major programming paradigms include: Imperative programming (how to perform tasks step by step), Declarative programming (what results you want), Object-Oriented Programming (organizing code and data into objects), Functional Programming (using pure functions without side effects), and Procedural Programming (based on procedure calls)."
                }
            ]
        }
        
        # Check if query matches any topic
        for topic, results in topics.items():
            if topic in query_lower:
                return results
                
        # No pre-defined results found
        return None

    def fetch_webpage(self, url):
        """Fetch and parse content from a webpage with multiple fallback strategies"""
        try:
            # Make sure URL has scheme
            if not url.startswith(('http://', 'https://')):
                url = 'https://' + url
            
            # Log the attempt
            self.errors.append(f"Attempting to fetch content from: {url}")
            
            # Set up headers that mimic a browser
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.127 Safari/537.36",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8"
            }
            
            # For the demo, we'll return simulated content
            # In a real implementation, you would make an actual HTTP request
            title = f"Simulated content for: {url}"
            content = f"This is simulated content for {url} containing relevant information about the search query. This would be real content from the web in a production environment."
            
            return {
                "url": url,
                "title": title,
                "content": content
            }
                
        except Exception as e:
            error_msg = f"Error fetching {url}: {str(e)}"
            self.errors.append(error_msg)
            return {
                "url": url,
                "title": "Error",
                "content": error_msg
            }

    def process_web_content(self, query):
        """Process web search results and create a vector store"""
        # Search the web
        search_results = self.web_search(query)
        
        # Track sources from the beginning
        self.sources = []
        for result in search_results:
            self.sources.append({
                "url": result["url"],
                "title": result["title"],
                "status": "Searched"
            })
        
        # Fetch and process documents
        documents = []
        for i, result in enumerate(search_results):
            doc = self.fetch_webpage(result["url"])
            documents.append(doc)
            
            # Update source status
            for source in self.sources:
                if source["url"] == result["url"]:
                    if "Error" in doc["title"]:
                        source["status"] = "Failed to retrieve"
                    else:
                        source["status"] = "Retrieved"
        
        # Set up vector store
        if documents:
            texts = []
            metadatas = []
            
            for doc in documents:
                chunks = self.text_splitter.split_text(doc["content"])
                for chunk in chunks:
                    texts.append(chunk)
                    metadatas.append({"source": doc["url"], "title": doc["title"]})
            
            # Create vector store
            self.web_vector_store = Chroma.from_texts(
                texts=texts,
                embedding=self.embeddings,
                metadatas=metadatas
            )
            
            return True
        return False

    def direct_retrieval_answer(self, query, user_id=None, mongodb=None, notebook_id=None):
        """Generate an answer using direct document retrieval.
        Requires documents to be processed first.
        """
        # Check if we have a vector store
        if not hasattr(self, 'doc_vector_store') or not self.doc_vector_store:
            return "Please upload and process documents first."
        
        try:
            # Get relevant chunks (safely with error handling)
            try:
                docs = self.doc_vector_store.similarity_search(query, k=4)
            except Exception as e:
                print(f"Error in similarity search: {str(e)}")
                docs = []
            
            # If no documents were retrieved, handle gracefully
            if not docs:
                return "No relevant information found in your documents."
            
            # Extract content from top chunks
            source_content = []
            for doc in docs:
                source_data = {
                    "content": doc.page_content,
                    "source": doc.metadata.get("source", "Unknown"),
                    "file_type": doc.metadata.get("file_type", "document")
                }
                
                # Add blockchain verification if available
                if "blockchain_verification" in doc.metadata:
                    source_data["blockchain_verification"] = doc.metadata["blockchain_verification"]
                    
                source_content.append(source_data)
            
            # Generate the answer
            chunks_text = ' '.join([doc.page_content for doc in docs])
            prompt = f"""
            Answer the following question based on the provided context from documents:
            
            Question: {query}
            
            Context:
            {chunks_text}
            
            Provide a clear, concise answer that addresses the question directly based on the information in the context.
            """
            
            answer = self.llm(prompt)
            
            # Log query to blockchain if enabled
            blockchain_log = None
            if self.use_blockchain and self.blockchain:
                try:
                    blockchain_log = self.log_query_blockchain(query, answer)
                except Exception as e:
                    print(f"Error logging to blockchain: {str(e)}")
            
            # Return result
            return {
                "answer": answer,
                "sources": source_content,
                "blockchain_log": blockchain_log
            }
        except Exception as e:
            error_msg = f"Error processing query: {str(e)}"
            print(error_msg)
            import traceback
            traceback.print_exc()
            return error_msg

    def enhanced_rag_answer(self, query, user_id=None, mongodb=None, notebook_id=None):
        """Generate an answer using the enhanced RAG pipeline.
        Will work with or without documents.
        """
        try:
            sources = []
            
            # Check if we have a vector store available
            has_documents = hasattr(self, 'doc_vector_store') and self.doc_vector_store
            
            if has_documents:
                # Get relevant document chunks (with error handling)
                try:
                    docs = self.doc_vector_store.similarity_search(query, k=4)
                except Exception as e:
                    print(f"Error in similarity search: {str(e)}")
                    docs = []
                
                if not docs:
                    # Fall back to simulated results if no documents are found
                    return self.hybrid_answer(query, user_id, mongodb, notebook_id)
                
                # Extract content from documents
                source_content = []
                for doc in docs:
                    source_data = {
                        "content": doc.page_content,
                        "source": doc.metadata.get("source", "Unknown"),
                        "file_type": doc.metadata.get("file_type", "document")
                    }
                    
                    # Add blockchain verification if available
                    if "blockchain_verification" in doc.metadata:
                        source_data["blockchain_verification"] = doc.metadata["blockchain_verification"]
                    
                    # Add to source list
                    sources.append(source_data)
                    source_content.append(doc.page_content)
                
                # Get initial answer based on document content
                chunks_text = ' '.join([doc.page_content for doc in docs])
                prompt = f"""
                Based on the following context, answer the question: {query}
                
                Context:
                {chunks_text}
                
                Answer:
                """
                
                initial_answer = self.llm(prompt)
                
                # Enhance the answer
                enhanced_answer = self.enhance_answer(initial_answer, query, source_content)
                
                # Log query to blockchain if enabled
                blockchain_log = None
                if self.use_blockchain and self.blockchain:
                    try:
                        blockchain_log = self.log_query_blockchain(query, enhanced_answer)
                    except Exception as e:
                        print(f"Error logging to blockchain: {str(e)}")
                
                return {
                    "answer": enhanced_answer,
                    "initial_answer": initial_answer,
                    "sources": sources,
                    "blockchain_log": blockchain_log
                }
            
            else:
                # No documents available, use llm directly with web research simulation
                simulated_results = self.simulate_search(query, num_results=3)
                web_sources = []
                
                # Process simulated web results
                context = ""
                for result in simulated_results:
                    context += f"Source: {result['title']}\nContent: {result['content']}\n\n"
                    
                    # Add to sources list
                    web_sources.append({
                        "content": result['content'],
                        "source": result['title'],
                        "file_type": "web" 
                    })
                
                # Generate a more comprehensive answer using web sources
                prompt = f"""
                I need to answer the following question: {query}
                
                I found these relevant information on the web:
                
                {context}
                
                Based on this information, provide a comprehensive answer to the question.
                """
                
                initial_answer = self.llm(prompt)
                
                # Generate enhanced answer
                enhanced_prompt = f"""
                I have an initial answer to the question: "{query}"
                
                Initial answer: {initial_answer}
                
                Please improve this answer by making it more comprehensive, accurate, and well-structured.
                Make sure to incorporate any relevant information from the web sources and organize the answer
                with clear explanations.
                """
                
                enhanced_answer = self.llm(enhanced_prompt)
                
                # Log query to blockchain if enabled
                blockchain_log = None
                if self.use_blockchain and self.blockchain:
                    try:
                        blockchain_log = self.log_query_blockchain(query, enhanced_answer)
                    except Exception as e:
                        print(f"Error logging to blockchain: {str(e)}")
                
                return {
                    "answer": enhanced_answer,
                    "initial_answer": initial_answer,
                    "sources": web_sources,
                    "blockchain_log": blockchain_log
                }
        except Exception as e:
            error_msg = f"Error processing query: {str(e)}"
            print(error_msg)
            import traceback
            traceback.print_exc()
            return error_msg

    def hybrid_answer(self, query, user_id=None, mongodb=None, notebook_id=None):
        """Generate an answer using hybrid search combining documents and web search.
        Will work even if no documents are uploaded.
        """
        # Initialize sources and content tracking
        doc_sources = []
        web_sources = []
        combined_context = ""
        
        # Check if we have a vector store available
        has_documents = hasattr(self, 'doc_vector_store') and self.doc_vector_store
        
        # Get document sources if available
        if has_documents:
            # Retrieve from document store
            try:
                docs = self.doc_vector_store.similarity_search(query, k=3)
                
                # Extract content from documents
                for doc in docs:
                    doc_content = doc.page_content
                    doc_source = doc.metadata.get("source", "Unknown document")
                    
                    # Add to combined context
                    combined_context += f"Document: {doc_source}\nContent: {doc_content}\n\n"
                    
                    # Create source data
                    source_data = {
                        "content": doc_content,
                        "source": doc_source,
                        "file_type": doc.metadata.get("file_type", "document")
                    }
                    
                    # Add blockchain verification if available
                    if "blockchain_verification" in doc.metadata:
                        source_data["blockchain_verification"] = doc.metadata["blockchain_verification"]
                    
                    # Add to sources
                    doc_sources.append(source_data)
            except Exception as e:
                print(f"Error retrieving documents: {str(e)}")
        
        # Always get web sources (simulated)
        web_results = self.simulate_search(query, num_results=3)
        
        # Process web results
        for result in web_results:
            # Add to combined context
            combined_context += f"Web: {result['title']}\nContent: {result['content']}\n\n"
            
            # Add to sources
            web_sources.append({
                "content": result['content'],
                "source": result['title'],
                "file_type": "web"
            })
        
        # All sources combined
        all_sources = doc_sources + web_sources
        
        # Generate comprehensive answer
        prompt = f"""
        I need to answer the following question thoroughly: {query}
        
        I have collected the following information:
        
        {combined_context}
        
        Based on all this information, provide a comprehensive, well-structured answer. 
        Integrate information from both documents and web sources when available.
        """
        
        answer = self.llm(prompt)
        
        # Log query to blockchain if enabled
        blockchain_log = None
        if self.use_blockchain and self.blockchain:
            try:
                blockchain_log = self.log_query_blockchain(query, answer)
            except Exception as e:
                print(f"Error logging to blockchain: {str(e)}")
        
        # Return combined result
        return {
            "answer": answer,
            "sources": all_sources,
            "doc_sources_count": len(doc_sources),
            "web_sources_count": len(web_sources),
            "blockchain_log": blockchain_log
        }

    def ask(self, query, mode="direct_retrieval", user_id=None, mongodb=None, notebook_id=None):
        """Ask a question and get an answer from the RAG system.
        
        Args:
            query: User question string
            mode: RAG mode to use (direct_retrieval, enhanced_rag, hybrid)
            user_id: Optional user ID for logging
            mongodb: Optional MongoDB connection for logging
            notebook_id: Optional notebook ID for context
            
        Returns:
            Answer dict or string
        """
        # Try to load vector store if not available
        if not hasattr(self, 'doc_vector_store') or not self.doc_vector_store:
            if mongodb and notebook_id:
                with st.spinner("Loading vector index..."):
                    # Try to load the FAISS index
                    load_success = self.load_vector_store(mongodb, notebook_id)
                    if not load_success:
                        # Try alternate loading method
                        self.load_faiss_only(mongodb, notebook_id)
        
        # Start timing for performance tracking
        start_time = time.time()
        
        # Log query if mongodb provided
        if mongodb and user_id:
            try:
                mongodb.log_query(user_id, query, 0, notebook_id)  # Time will be updated later
            except Exception as e:
                print(f"Error logging query to MongoDB: {str(e)}")
        
        try:
            # Call the appropriate method based on mode
            if mode == "enhanced_rag":
                result = self.enhanced_rag_answer(query, user_id, mongodb, notebook_id)
            elif mode == "hybrid":
                result = self.hybrid_answer(query, user_id, mongodb, notebook_id)
            else:  # default to direct_retrieval
                # Only direct retrieval requires documents
                if not hasattr(self, 'doc_vector_store') or not self.doc_vector_store:
                    return "Please upload and process documents first."
                result = self.direct_retrieval_answer(query, user_id, mongodb, notebook_id)
                
            # Calculate query time
            query_time = time.time() - start_time
            
            # Update response with query time and mode
            if isinstance(result, dict):
                result["query_time"] = query_time
                result["mode"] = mode
            
            # Update log with actual time if mongodb provided
            if mongodb and user_id:
                try:
                    mongodb.log_query(user_id, query, query_time, notebook_id)
                except Exception as e:
                    print(f"Error updating query log in MongoDB: {str(e)}")
                    
            return result
            
        except Exception as e:
            # Handle exceptions gracefully with detailed error messages
            error_message = f"Error processing query: {str(e)}"
            print(error_message)
            
            # Print full stack trace for debugging
            import traceback
            traceback.print_exc()
            
            return f"Error processing query: {str(e)}"

    def detect_query_domain(self, query):
        """Detect potential domains from a query.
        
        Args:
            query: User question string
            
        Returns:
            List of potential domain strings in priority order
        """
        # Simple keyword matching for now
        domain_keywords = {
            "machine learning": ["machine learning", "ml", "algorithm", "model", "training", "neural"],
            "data science": ["data science", "analysis", "statistics", "visualization"],
            "programming": ["programming", "code", "function", "class", "variable"],
            "finance": ["finance", "market", "investment", "stock", "financial"],
            "healthcare": ["healthcare", "medical", "patient", "treatment", "clinical"]
        }
        
        query_lower = query.lower()
        matches = []
        
        # First check for exact domain mentions
        for domain, keywords in domain_keywords.items():
            if domain.lower() in query_lower:
                matches.append(domain)
        
        # Then check for keyword matches
        for domain, keywords in domain_keywords.items():
            if domain not in matches:  # Skip if already matched
                for keyword in keywords:
                    if keyword.lower() in query_lower:
                        matches.append(domain)
                        break
        
        return matches

    def load_faiss_index(self, notebook_id=None, user_id=None, domain=None, mongodb=None):
        """Load a previously saved FAISS index by notebook ID or domain.
        
        Args:
            notebook_id: ID of the notebook (optional if using domain)
            user_id: User ID for domain search (required if using domain)
            domain: Domain/topic to search for (optional if using notebook_id)
            mongodb: MongoDB connection
            
        Returns:
            Boolean indicating success
        """
        if not mongodb:
            return False
            
        try:
            import faiss
            import pickle
            import base64
            from io import BytesIO
            from langchain_community.vectorstores import FAISS
            
            # Get the saved index from MongoDB
            success, result = mongodb.get_faiss_index(notebook_id)
            
            if not success:
                return False
                
            # Extract data
            faiss_index = result["faiss_index"]
            documents = result["documents"]
            
            # Deserialize FAISS index
            index_bytes = base64.b64decode(faiss_index)
            index_buffer = BytesIO(index_bytes)
            faiss_index = faiss.read_index(index_buffer)
            
            # Deserialize documents
            docs_bytes = base64.b64decode(documents)
            docs_buffer = BytesIO(docs_bytes)
            documents = pickle.load(docs_buffer)
            
            # Store documents for later use
            self.documents = documents
            
            # Create vector store with the loaded index
            embedding_function = self.embeddings
            self.doc_vector_store = FAISS(embedding_function, faiss_index, documents, {})
            
            # Store notebook_id if we loaded by domain
            if domain and not notebook_id:
                self.current_notebook_id = result.get("notebook_id")
                
            return True
            
        except Exception as e:
            if hasattr(self, '_suppress_messages'):
                self._suppress_messages = True  # Suppress duplicate error messages
            st.error(f"Error loading vector index: {str(e)}")
            if hasattr(self, '_suppress_messages'):
                self._suppress_messages = False
            return False

    def get_performance_metrics(self):
        """Return performance metrics for the RAG system."""
        if not self.processing_times:
            return None
            
        return {
            "documents_processed": self.documents_processed,
            "index_building_time": self.processing_times.get("index_building", 0),
            "total_processing_time": self.processing_times.get("total_time", 0),
            "memory_used_gb": self.processing_times.get("memory_used_gb", 0),
            "device": self.device,
            "embedding_model": self.embedding_model_name,
            "errors": len(self.errors),
            "blockchain_enabled": hasattr(self, 'use_blockchain') and self.use_blockchain
        }

    def save_vector_store(self, mongodb, notebook_id, user_id, is_nested=False):
        """Save the current vector store to MongoDB for the notebook.
        
        Args:
            mongodb: MongoDB connection
            notebook_id: ID of the notebook
            user_id: User ID
            is_nested: Whether this is being called from another streamlit component
        
        Returns:
            Boolean indicating success
        """
        try:
            import faiss
            import pickle
            import tempfile
            import os
            
            # Make sure we have a vector store to save
            if not hasattr(self, 'doc_vector_store') or not self.doc_vector_store:
                raise ValueError("No vector store available to save")
            
            # 1. First save the FAISS index to a temporary file
            temp_dir = tempfile.mkdtemp()
            index_path = os.path.join(temp_dir, 'faiss_index.bin')
            
            # Get the index from the vector store
            faiss_index = self.doc_vector_store.index
            
            # Write to temporary file
            faiss.write_index(faiss_index, index_path)
            
            # Read binary data
            with open(index_path, 'rb') as f:
                index_binary = f.read()
            
            # 2. Get document info directly from the doc_vector_store
            # LangChain FAISS stores documents in the docstore attribute
            if hasattr(self.doc_vector_store, 'docstore'):
                docs_dict = self.doc_vector_store.docstore._dict
            # Or access via the _dict attribute (older versions)
            elif hasattr(self.doc_vector_store, '_dict'):
                docs_dict = self.doc_vector_store._dict
            else:
                raise ValueError("Cannot find documents in vector store")
            
            # 3. Serialize the documents dictionary
            doc_path = os.path.join(temp_dir, 'documents.pkl')
            with open(doc_path, 'wb') as f:
                pickle.dump(docs_dict, f)
            
            # Read binary data
            with open(doc_path, 'rb') as f:
                documents_binary = f.read()
            
            # 4. Count documents in vector store
            try:
                # Try to get document count from docstore
                doc_count = len(docs_dict)
            except:
                # Fallback count
                doc_count = self.doc_vector_store.index.ntotal
            
            # 5. Prepare metadata
            metadata = {
                "embedding_model": self.embedding_model_name,
                "chunk_size": self.chunk_size,
                "chunk_overlap": self.chunk_overlap,
                "document_count": doc_count,
                "index_size_bytes": len(index_binary),
                "blockchain_enabled": hasattr(self, 'use_blockchain') and self.use_blockchain
            }
            
            # 6. Save to MongoDB
            success, message = mongodb.save_faiss_index(
                notebook_id, user_id, index_binary, documents_binary, metadata
            )
            
            # 7. Clean up temp files
            try:
                os.remove(index_path)
                os.remove(doc_path)
                os.rmdir(temp_dir)
            except:
                pass
            
            # 8. Show message if successful
            if success and not is_nested:
                st.success("Vector index saved successfully!")
            
            return success
            
        except Exception as e:
            error_msg = f"Error creating vector store: {str(e)}"
            print(error_msg)  # Log to console for debugging
            if not is_nested:
                st.error(error_msg)
            return False
    
    def load_vector_store(self, mongodb, notebook_id):
        """Load a vector store from MongoDB for the notebook.
        
        Args:
            mongodb: MongoDB connection
            notebook_id: ID of the notebook
        
        Returns:
            Boolean indicating success
        """
        try:
            import faiss
            import pickle
            import tempfile
            import os
            from langchain_community.vectorstores import FAISS
            
            # 1. Get data from MongoDB
            success, result = mongodb.get_faiss_index(notebook_id)
            
            if not success:
                print(f"No vector index found for notebook {notebook_id}")
                return False
            
            # Check if result contains required data
            if not result or "faiss_index" not in result or "documents" not in result:
                print("Missing required data in retrieved index")
                return False
                
            # 2. Create temporary files
            temp_dir = tempfile.mkdtemp()
            index_path = os.path.join(temp_dir, 'faiss_index.bin')
            doc_path = os.path.join(temp_dir, 'documents.pkl')
            
            # 3. Write binary data to temp files
            try:
                with open(index_path, 'wb') as f:
                    f.write(result["faiss_index"])
                
                with open(doc_path, 'wb') as f:
                    f.write(result["documents"])
            except Exception as e:
                print(f"Error writing temporary files: {str(e)}")
                return False
            
            # 4. Load FAISS index from file
            try:
                faiss_index = faiss.read_index(index_path)
            except Exception as e:
                print(f"Error reading FAISS index: {str(e)}")
                os.remove(index_path)
                os.remove(doc_path)
                os.rmdir(temp_dir)
                return False
            
            # 5. Load document dictionary from pickle file
            try:
                with open(doc_path, 'rb') as f:
                    docs_dict = pickle.load(f)
            except Exception as e:
                print(f"Error loading document dictionary: {str(e)}")
                os.remove(index_path)
                os.remove(doc_path)
                os.rmdir(temp_dir)
                return False
            
            # Store documents for later use
            self.documents = list(docs_dict.values())
            
            # 6. Create vector store from index and documents
            try:
                self.doc_vector_store = FAISS(
                    self.embedding_function,
                    faiss_index,
                    {},  # Empty mapping since we're directly using the docstore
                    {}   # Empty retriever options
                )
            except Exception as e:
                print(f"Error creating vector store: {str(e)}")
                os.remove(index_path)
                os.remove(doc_path)
                os.rmdir(temp_dir)
                return False
            
            # 7. Set the docstore manually
            try:
                from langchain.docstore.document import DocumentStore
                docstore = DocumentStore()
                docstore._dict = docs_dict
                self.doc_vector_store.docstore = docstore
            except Exception as e:
                print(f"Error setting docstore: {str(e)}")
                self.doc_vector_store = None  # Reset if we failed
                os.remove(index_path)
                os.remove(doc_path)
                os.rmdir(temp_dir)
                return False
            
            # 8. Clean up temp files
            try:
                os.remove(index_path)
                os.remove(doc_path)
                os.rmdir(temp_dir)
            except:
                pass
            
            # Return success
            count = len(docs_dict) if hasattr(docs_dict, '__len__') else "unknown number of"
            st.success(f"Loaded {count} documents from saved vectors!")
            return True
            
        except Exception as e:
            error_msg = f"Error loading vector store: {str(e)}"
            print(error_msg)  # Log to console for debugging
            import traceback
            traceback.print_exc()
            st.error(error_msg)
            return False

    def debug_vector_store(self):
        """Debug utility to inspect vector store structure.
        Prints out all available attributes and methods.
        """
        if not hasattr(self, 'doc_vector_store') or not self.doc_vector_store:
            print("No vector store available!")
            return
        
        print("Vector store type:", type(self.doc_vector_store))
        print("\nVector store attributes:")
        for attr in dir(self.doc_vector_store):
            if not attr.startswith('_'):
                try:
                    value = getattr(self.doc_vector_store, attr)
                    print(f"- {attr}: {type(value)}")
                except:
                    print(f"- {attr}: <error accessing>")
        
        # Check if the vector store has serialization methods
        if hasattr(self.doc_vector_store, 'serialize_to_bytes'):
            print("\nVector store has serialize_to_bytes method")
        
        # Check for save and load methods
        if hasattr(self.doc_vector_store, 'save_local'):
            print("Vector store has save_local method")
        
        # Print all attributes of self (EnhancedRAG)
        print("\nEnhancedRAG attributes:")
        for attr in dir(self):
            if not attr.startswith('_') and attr != 'debug_vector_store':
                try:
                    value = getattr(self, attr)
                    print(f"- {attr}: {type(value)}")
                except:
                    print(f"- {attr}: <error accessing>")

    def load_faiss_only(self, mongodb, notebook_id):
        """Load just the FAISS index without relying on documents.
        This will recreate a vector store with just the index.
        """
        try:
            import faiss
            import tempfile
            import os
            from langchain_community.vectorstores import FAISS
            
            # 1. Get data from MongoDB
            success, result = mongodb.get_faiss_index(notebook_id)
            
            if not success:
                print(f"No vector index found for notebook {notebook_id}")
                return False
            
            # Check if result contains required data
            if "faiss_index" not in result:
                print("Missing FAISS index data in retrieved result")
                return False
                
            # 2. Create temporary file for the index
            temp_dir = tempfile.mkdtemp()
            index_path = os.path.join(temp_dir, 'faiss_index.bin')
            
            # 3. Write the index data to file
            try:
                with open(index_path, 'wb') as f:
                    f.write(result["faiss_index"])
            except Exception as e:
                print(f"Error writing FAISS index to temp file: {str(e)}")
                os.rmdir(temp_dir)
                return False
            
            # 4. Load the FAISS index
            try:
                faiss_index = faiss.read_index(index_path)
            except Exception as e:
                print(f"Error reading FAISS index: {str(e)}")
                os.remove(index_path)
                os.rmdir(temp_dir)
                return False
            
            # 5. Create a vector store with just the index
            try:
                self.doc_vector_store = FAISS(
                    self.embedding_function,
                    faiss_index,
                    {},
                    {}
                )
            except Exception as e:
                print(f"Error creating vector store: {str(e)}")
                os.remove(index_path)
                os.rmdir(temp_dir)
                return False
            
            # 6. Clean up
            try:
                os.remove(index_path)
                os.rmdir(temp_dir)
            except:
                pass
            
            # Success message
            st.success("Loaded vector index for similarity search")
            return True
            
        except Exception as e:
            error_msg = f"Error loading vector store: {str(e)}"
            print(error_msg)
            import traceback
            traceback.print_exc()
            st.error(error_msg)
            return False